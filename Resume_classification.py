import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
import joblib as jb
import tempfile, subprocess, os ,re
from wordcloud import WordCloud
from pdf2docx import Converter
from docx import Document
from nltk.stem import WordNetLemmatizer
from nltk.corpus import stopwords

# doc to docx
def convert_doc_to_docx(input_path, output_path):
    """
    Converts .doc to .docx using LibreOffice in headless mode.
    """
    # libreoffice path 
    libreoffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
    # make sure output path exists
    os.makedirs(output_path, exist_ok=True)
    subprocess.run([
        libreoffice_path, "--headless", "--convert-to", "docx", "--outdir",
        output_path ,input_path
    ], check=True)
    return os.path.join(output_path, os.path.splitext(os.path.basename(input_path))[0] + ".docx")

# pdf to docx
def convert_pdf_to_docx(input_path, output_path):
    """
    Converts .pdf to .docx using pdf2docx.
    """
    cv = Converter(input_path)
    # Read full pdf
    cv.convert(output_path, start=0, end=None)
    cv.close()

 
def handle_uploaded_file(uploaded_file):
    """
    Takes uploaded file (.doc, .pdf, .docx) from Streamlit and returns the final .docx path.
    """
    # Save uploaded file temporarily
    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmp_file:
        tmp_file.write(uploaded_file.read())
        tmp_path = tmp_file.name

    if uploaded_file.name.lower().endswith(".doc"):
        output_dir = os.path.dirname(tmp_path)  # use same folder as temp file
        tmp_docx_path = convert_doc_to_docx(tmp_path, output_dir)

    elif uploaded_file.name.lower().endswith(".pdf"):
        tmp_docx_path = tmp_path + ".docx"
        convert_pdf_to_docx(tmp_path, tmp_docx_path)

    elif uploaded_file.name.lower().endswith(".docx"):
        tmp_docx_path = tmp_path

    else:
        raise ValueError("Unsupported file format!")

    return tmp_docx_path

lemma = WordNetLemmatizer()
stop_words = set(stopwords.words('english'))

known_skills = {
    "sql", "ps", "wf", "react", "js", "wd", "api", "db", "etl", "ai",
    "pl/sql", "html", "css", "aws", "json", "rest", "soap", "ci/cd",
}

def clean_text(text):
    # Remove non-letters and lowercase
    text = re.sub('[^a-zA-Z]', ' ', text) # keep only alphabets
    text = text.lower() # convert to lowercase
    # Tokenize
    words = text.split()
    
    # Lemmatize words except known skills, remove stopwords
    cleaned_words = [
        word if word in known_skills else lemma.lemmatize(word)
        for word in words
        if word not in stop_words
    ]
    return ' '.join(cleaned_words)


def extract_full_text(file_path):
    doc = Document(file_path)
    full_text = []

    # Extract paragraph text
    for para in doc.paragraphs:
        if para.text.strip(): 
            full_text.append(para.text.strip()) # remove spaces and append to list

    # Extract table text
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                 # join '|' to differenciate between table text and normal text (optional)
                full_text.append(' | '.join(row_data))

    return '\n'.join(full_text)

# load model
model = jb.load("C:\\Users\\gsree\\Deployment1\\Resume_Pipeline1.pkl")

# creating 2 tabs for prediction and wordcloud
tab0, tab1 = st.tabs(['Resume Classifier', 'WordCloud'])
with tab0:
    # Streamlit section
    st.title("Upload & Convert to DOCX")
    st.info(
        "Note: The model may misclassify some resumes because it is trained on only four categories. "
        "It can also be influenced by common words that appear across different job roles."
    )
    # labels in sequence as created by label encoder 
    classes = ['Peoplesoft Resume', 'React Developer', 'SQL Developer', 'workday']
    uploaded_file = st.file_uploader("Upload a DOC, DOCX, or PDF", type=["doc", "docx", "pdf"])

    if uploaded_file:
        # Save the state to use in tab1
        st.session_state.uploaded_file = uploaded_file
        try:
            docx_path = handle_uploaded_file(uploaded_file)
            if not uploaded_file.name.lower().endswith(".docx"):
                st.success(f"File converted to DOCX")
            resume_text = extract_full_text(docx_path)
            st.text_area("Extracted Resume Text (conversion text)", resume_text[0::],height=600)

            cleaned_text = clean_text(resume_text)
            X = model.named_steps['tfidf'].transform([cleaned_text])
            # check if X does not contains zero 
            if X.nnz != 0:
                # Get predicted probabilities
                probs = model.predict_proba([cleaned_text])[0]
                # Get max probability 
                predicted_index = probs.argmax()
                predicted_label = classes[predicted_index]
                confidence = probs[predicted_index]

                # Confidence threshold
                THRESHOLD = 0.4  # tune as needed
                if confidence < THRESHOLD:
                    st.warning("No known category detected (confidence too low).")
                    final_prediction = None
                else:
                    final_prediction = predicted_label
                    st.write(f'Predicted Class: {final_prediction} (Confidence: {confidence:.2f})')
            else:
                st.warning("No meaningful text features found to predict.")  
        except Exception as e:
            st.error(f"Error: {e}")

with tab1:
    # check if file exists or not 
    if 'uploaded_file' in st.session_state:
        if st.button('Generate WordCloud'):
            if cleaned_text.strip():  # Check that text is not empty after cleaning
                st.markdown('**WordCloud for User Input**')
                wordcloud_user = WordCloud(width=1000, height=600, background_color='black',
                                       colormap='Pastel1').generate(cleaned_text)
                fig1, ax1 = plt.subplots(figsize=(10, 5))
                ax1.imshow(wordcloud_user)
                ax1.axis('off')
                st.pyplot(fig1)
            else:
                st.warning("No meaningful words left after cleaning. Try entering more descriptive text.")
    else:
        st.info("Please upload a file in Tab 0 first")